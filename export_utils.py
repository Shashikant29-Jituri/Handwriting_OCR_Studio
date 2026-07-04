"""
export_utils.py
Generates .docx (python-docx) and .pdf (ReportLab) exports from OCR results,
preserving basic layout (headings, bullets, tables) and rendering the correct
font per script (Latin / Devanagari / Kannada / Tamil / Telugu / Bengali).
"""

import io
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from ocr_utils import LANGUAGE_FONTS

# ---------------------------------------------------------------------------
# Font files
# Place actual .ttf files under ./fonts/ (Noto Sans family covers all scripts).
# Download once, e.g.:
#   NotoSans-Regular.ttf, NotoSansDevanagari-Regular.ttf, NotoSansKannada-Regular.ttf,
#   NotoSansTamil-Regular.ttf, NotoSansTelugu-Regular.ttf, NotoSansBengali-Regular.ttf
# from https://fonts.google.com/noto
# ---------------------------------------------------------------------------

FONT_DIR = os.path.join(os.path.dirname(__file__), "fonts")

FONT_FILES = {
    "NotoSans": "NotoSans-Regular.ttf",
    "NotoSansDevanagari": "NotoSansDevanagari-Regular.ttf",
    "NotoSansKannada": "NotoSansKannada-Regular.ttf",
    "NotoSansTamil": "NotoSansTamil-Regular.ttf",
    "NotoSansTelugu": "NotoSansTelugu-Regular.ttf",
    "NotoSansBengali": "NotoSansBengali-Regular.ttf",
}

_registered_fonts = set()


def _register_pdf_font(font_key: str) -> str:
    """Registers a TTF font with ReportLab, falls back to Helvetica if the file is missing."""
    if font_key in _registered_fonts:
        return font_key
    font_path = os.path.join(FONT_DIR, FONT_FILES.get(font_key, ""))
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont(font_key, font_path))
        _registered_fonts.add(font_key)
        return font_key
    return "Helvetica"  # fallback -- won't render non-Latin scripts correctly


def _docx_set_run_font(run, font_name: str, size: int = 11):
    run.font.size = Pt(size)
    run.font.name = font_name
    # Ensure the East-Asian/complex-script font attribute is also set so Word
    # picks the right font for non-Latin glyphs.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rfonts is None:
        from docx.oxml.ns import qn
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    from docx.oxml.ns import qn
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:cs'), font_name)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def build_docx(results: list, combine_pages: bool = True) -> bytes:
    """
    results: list of ocr_utils.OcrResult (already run through OCR).
    Produces one .docx with a heading per file/page, preserving bullet/heading/table structure.
    """
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    for i, res in enumerate(results):
        lang_code = None
        from ocr_utils import LANGUAGES
        lang_code = LANGUAGES.get(res.language, "en")
        font_key = LANGUAGE_FONTS.get(lang_code, "NotoSans")

        title = doc.add_heading(f"{res.filename} (page {res.page_number}) -- {res.language}", level=1)

        if res.error:
            p = doc.add_paragraph(f"[Error extracting this page: {res.error}]")
            p.runs[0].font.color.rgb = None
            continue

        for block in res.blocks:
            if block.kind == "heading":
                p = doc.add_paragraph()
                run = p.add_run(block.text)
                run.bold = True
                _docx_set_run_font(run, font_key, size=14 - block.level)
            elif block.kind == "bullet":
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(block.text)
                _docx_set_run_font(run, font_key)
            elif block.kind == "table_row":
                # Group consecutive table_row blocks into one table
                pass  # handled in the pre-pass below
            else:
                p = doc.add_paragraph()
                run = p.add_run(block.text)
                _docx_set_run_font(run, font_key)

        # Second pass: render tables as actual Word tables (grouping consecutive rows)
        table_rows = [b for b in res.blocks if b.kind == "table_row"]
        if table_rows:
            n_cols = max(len(r.cells) for r in table_rows)
            table = doc.add_table(rows=0, cols=n_cols)
            table.style = "Table Grid"
            for row_block in table_rows:
                row_cells = table.add_row().cells
                for c_idx in range(n_cols):
                    text = row_block.cells[c_idx] if c_idx < len(row_block.cells) else ""
                    run = row_cells[c_idx].paragraphs[0].add_run(text)
                    _docx_set_run_font(run, font_key, size=10)

        if i != len(results) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def build_pdf(results: list) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=20 * mm, rightMargin=20 * mm,
                             topMargin=18 * mm, bottomMargin=18 * mm)
    story = []

    from ocr_utils import LANGUAGES

    for i, res in enumerate(results):
        lang_code = LANGUAGES.get(res.language, "en")
        font_key = _register_pdf_font(LANGUAGE_FONTS.get(lang_code, "NotoSans"))

        heading_style = ParagraphStyle("Heading", fontName=font_key, fontSize=14,
                                        leading=18, spaceAfter=10, textColor=colors.HexColor("#1a1a1a"))
        body_style = ParagraphStyle("Body", fontName=font_key, fontSize=11, leading=16, spaceAfter=6)
        bullet_style = ParagraphStyle("Bullet", fontName=font_key, fontSize=11, leading=16,
                                       leftIndent=14, spaceAfter=4, bulletIndent=4)

        story.append(Paragraph(f"{res.filename} (page {res.page_number}) -- {res.language}", heading_style))
        story.append(Spacer(1, 6))

        if res.error:
            story.append(Paragraph(f"[Error extracting this page: {res.error}]", body_style))
            if i != len(results) - 1:
                story.append(PageBreak())
            continue

        table_rows_buffer = []

        def flush_table():
            if not table_rows_buffer:
                return
            data = [row.cells for row in table_rows_buffer]
            n_cols = max(len(r) for r in data)
            data = [r + [""] * (n_cols - len(r)) for r in data]
            t = Table(data, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("FONT", (0, 0), (-1, -1), font_key, 10),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))
            table_rows_buffer.clear()

        for block in res.blocks:
            if block.kind == "table_row":
                table_rows_buffer.append(block)
                continue
            else:
                flush_table()

            if block.kind == "heading":
                story.append(Paragraph(block.text, heading_style))
            elif block.kind == "bullet":
                story.append(Paragraph(f"&bull;&nbsp;&nbsp;{block.text}", bullet_style))
            else:
                story.append(Paragraph(block.text.replace("\n", "<br/>"), body_style))

        flush_table()

        if i != len(results) - 1:
            story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()
